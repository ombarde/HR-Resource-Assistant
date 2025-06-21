from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import json
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
from datetime import datetime
from employees_data import EMPLOYEES_DATA
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
import io

app = FastAPI(title="HR Resource Query Chatbot", version="1.0.0")

# Debug: Print all registered routes at startup
import sys
if __name__ == "__main__" or "uvicorn" in sys.argv[0]:
    print("Registered routes:")
    for route in app.routes:
        print(f"{route.path} [{','.join(route.methods)}]")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ChatQuery(BaseModel):
    query: str

class ChatResponse(BaseModel):
    response: str
    matched_employees: List[Dict[str, Any]]
    confidence: float

class EmployeeSearchQuery(BaseModel):
    skills: Optional[List[str]] = None
    min_experience: Optional[int] = None
    availability: Optional[str] = None
    department: Optional[str] = None

# RAG System Implementation
class HRChatbot:
    def __init__(self):
        self.employees = EMPLOYEES_DATA["employees"]
        self.vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
        self._prepare_embeddings()
    
    def _prepare_embeddings(self):
        """Prepare employee text representations and embeddings"""
        self.employee_texts = []
        for emp in self.employees:
            # Create comprehensive text representation
            text = f"{emp['name']} {emp['role']} {' '.join(emp['skills'])} {' '.join(emp['projects'])} {emp['department']} {emp['availability']} {emp['experience_years']} years experience"
            self.employee_texts.append(text.lower())
        
        # Create TF-IDF embeddings
        self.embeddings = self.vectorizer.fit_transform(self.employee_texts)
    
    def _extract_query_features(self, query: str) -> Dict[str, Any]:
        """Extract features from natural language query"""
        query_lower = query.lower()
        features = {
            'skills': [],
            'experience': None,
            'availability': None,
            'project_domain': None,
            'department': None
        }
        
        # Extract skills
        common_skills = [
            'python', 'javascript', 'react', 'node.js', 'aws', 'docker', 'kubernetes',
            'java', 'spring boot', 'postgresql', 'mongodb', 'sql', 'tensorflow',
            'pytorch', 'machine learning', 'ml', 'ai', 'flutter', 'react native',
            'vue.js', 'angular', 'django', 'fastapi', 'express', 'redis', 'kafka',
            'go', 'rust', 'typescript', 'graphql', 'rest api', 'microservices'
        ]
        
        for skill in common_skills:
            if skill in query_lower:
                features['skills'].append(skill)
        
        # Extract experience requirements
        exp_match = re.search(r'(\d+)\+?\s*years?', query_lower)
        if exp_match:
            features['experience'] = int(exp_match.group(1))
        
        # Extract availability
        if 'available' in query_lower or 'free' in query_lower:
            features['availability'] = 'available'
        
        # Extract project domains
        domains = ['healthcare', 'finance', 'banking', 'e-commerce', 'mobile', 'web', 'analytics']
        for domain in domains:
            if domain in query_lower:
                features['project_domain'] = domain
        
        return features
    
    def search_employees(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search employees using semantic similarity"""
        # Create query embedding
        query_embedding = self.vectorizer.transform([query.lower()])
        
        # Calculate similarities
        similarities = cosine_similarity(query_embedding, self.embeddings).flatten()
        
        # Get top matches
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0.1:  # Minimum similarity threshold
                employee = self.employees[idx].copy()
                employee['similarity_score'] = float(similarities[idx])
                results.append(employee)
        
        return results
    
    def generate_response(self, query: str, matched_employees: List[Dict[str, Any]]) -> str:
        """Generate natural language response"""
        if not matched_employees:
            return "I couldn't find any employees matching your criteria. Could you please provide more specific requirements or try different keywords?"
        
        features = self._extract_query_features(query)
        
        # Start response
        response = f"Based on your requirements, I found {len(matched_employees)} suitable candidate{'s' if len(matched_employees) > 1 else ''}:\n\n"
        
        for i, emp in enumerate(matched_employees[:3], 1):  # Limit to top 3 for readability
            response += f"**{emp['name']}** would be {'an excellent' if i == 1 else 'a great'} choice. "
            response += f"{'She' if 'Dr.' in emp['name'] or emp['name'].split()[0] in ['Alice', 'Emily', 'Lisa', 'Maria', 'Jennifer', 'Amanda', 'Sophie'] else 'He'} has {emp['experience_years']} years of experience "
            
            # Highlight relevant skills
            relevant_skills = [skill for skill in emp['skills'] if any(s in skill.lower() for s in features['skills'])]
            if relevant_skills:
                response += f"with expertise in {', '.join(relevant_skills[:3])}. "
            else:
                response += f"and specializes in {', '.join(emp['skills'][:3])}. "
            
            # Mention relevant projects
            if features['project_domain']:
                relevant_projects = [proj for proj in emp['projects'] if features['project_domain'].lower() in proj.lower()]
                if relevant_projects:
                    response += f"{'She' if 'Dr.' in emp['name'] or emp['name'].split()[0] in ['Alice', 'Emily', 'Lisa', 'Maria', 'Jennifer', 'Amanda', 'Sophie'] else 'He'} worked on {relevant_projects[0]}, which is directly relevant to your needs. "
            else:
                response += f"Recent projects include {emp['projects'][0]}. "
            
            # Availability
            if emp['availability'] == 'available':
                response += "Currently available for new assignments.\n\n"
            else:
                response += "Currently busy but may be available soon.\n\n"
        
        if len(matched_employees) > 3:
            response += f"I found {len(matched_employees) - 3} additional candidates. Would you like me to provide more details or help you narrow down the requirements?"
        else:
            response += "Would you like me to provide more details about their specific experience or check their availability for meetings?"
        
        return response

# Initialize chatbot
chatbot = HRChatbot()

def score_employee_for_query(employee: Dict[str, Any], query: str) -> float:
    """Simple heuristic to score employee relevance to query"""
    query_lower = query.lower()
    score = 0.0

    # Score based on skill matches
    skills = [skill.lower() for skill in employee.get('skills', [])]
    for skill in skills:
        if skill in query_lower:
            score += 1.0

    # Score based on experience
    exp_match = re.search(r'(\d+)\+?\s*years?', query_lower)
    if exp_match:
        required_exp = int(exp_match.group(1))
        if employee.get('experience_years', 0) >= required_exp:
            score += 1.0

    # Score based on availability
    if 'available' in query_lower or 'free' in query_lower:
        if employee.get('availability', '').lower() == 'available':
            score += 1.0

    return score

@app.post("/recommendations")
async def get_recommendations(query: ChatQuery):
    """Return AI-powered candidate recommendations based on query"""
    try:
        scored_employees = []
        for emp in EMPLOYEES_DATA["employees"]:
            score = score_employee_for_query(emp, query.query)
            if score > 0:
                emp_copy = emp.copy()
                emp_copy['recommendation_score'] = score
                scored_employees.append(emp_copy)

        # Sort by score descending
        scored_employees.sort(key=lambda e: e['recommendation_score'], reverse=True)

        # Return top 5 recommendations
        top_recommendations = scored_employees[:5]

        return {"recommendations": top_recommendations, "count": len(top_recommendations)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating recommendations: {str(e)}")

@app.get("/")
async def root():
    return {"message": "HR Resource Query Chatbot API", "version": "1.0.0"}

@app.post("/chat", response_model=ChatResponse)
async def chat_query(query: ChatQuery):
    """Process natural language HR queries"""
    try:
        # Search for matching employees
        matched_employees = chatbot.search_employees(query.query)
        
        # Generate natural language response
        response_text = chatbot.generate_response(query.query, matched_employees)
        
        # Calculate overall confidence
        confidence = np.mean([emp['similarity_score'] for emp in matched_employees]) if matched_employees else 0.0
        
        return ChatResponse(
            response=response_text,
            matched_employees=matched_employees,
            confidence=confidence
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")

@app.get("/employees")
async def get_all_employees():
    """Get all employees"""
    return EMPLOYEES_DATA

@app.get("/employees/search")
async def search_employees(
    skills: Optional[str] = None,
    min_experience: Optional[int] = None,
    availability: Optional[str] = None,
    department: Optional[str] = None
):
    """Search employees with specific filters"""
    filtered_employees = EMPLOYEES_DATA["employees"]
    
    if skills:
        skill_list = [s.strip().lower() for s in skills.split(',')]
        filtered_employees = [
            emp for emp in filtered_employees
            if any(skill.lower() in [s.lower() for s in emp['skills']] for skill in skill_list)
        ]
    
    if min_experience:
        filtered_employees = [
            emp for emp in filtered_employees
            if emp['experience_years'] >= min_experience
        ]
    
    if availability:
        filtered_employees = [
            emp for emp in filtered_employees
            if emp['availability'].lower() == availability.lower()
        ]
    
    if department:
        filtered_employees = [
            emp for emp in filtered_employees
            if department.lower() in emp['department'].lower()
        ]
    
    return {"employees": filtered_employees, "count": len(filtered_employees)}

@app.get("/employees/{employee_id}")
async def get_employee(employee_id: int):
    """Get specific employee by ID"""
    employee = next((emp for emp in EMPLOYEES_DATA["employees"] if emp["id"] == employee_id), None)
    if not employee:
        raise HTTPException(status_code=404, detail="Employee not found")
    return employee

@app.get("/stats")
async def get_stats():
    """Get basic statistics about the employee database"""
    employees = EMPLOYEES_DATA["employees"]
    
    # Count by availability
    available_count = len([emp for emp in employees if emp['availability'] == 'available'])
    busy_count = len([emp for emp in employees if emp['availability'] == 'busy'])
    
    # Count by department
    departments = {}
    for emp in employees:
        dept = emp['department']
        departments[dept] = departments.get(dept, 0) + 1
    
    # Average experience
    avg_experience = np.mean([emp['experience_years'] for emp in employees])
    
    # Most common skills
    all_skills = []
    for emp in employees:
        all_skills.extend(emp['skills'])
    
    skill_counts = {}
    for skill in all_skills:
        skill_counts[skill] = skill_counts.get(skill, 0) + 1
    
    top_skills = sorted(skill_counts.items(), key=lambda x: x[1], reverse=True)[:10]
    
    return {
        "total_employees": len(employees),
        "available": available_count,
        "busy": busy_count,
        "departments": departments,
        "average_experience": round(avg_experience, 1),
        "top_skills": dict(top_skills)
    }

from fastapi import Body
import os
import json

EMPLOYEE_DATA_FILE = os.path.join(os.path.dirname(__file__), 'employees_data.json')

if os.path.exists(EMPLOYEE_DATA_FILE):
    with open(EMPLOYEE_DATA_FILE, 'r', encoding='utf-8') as f:
        EMPLOYEES_DATA = json.load(f)
else:
    from employees_data import EMPLOYEES_DATA

@app.post("/employees")
async def add_employee(employee: Dict = Body(...)):
    """Add a new employee to the dataset and save persistently"""
    required_fields = {"id", "name", "skills", "experience_years", "projects", "availability", "department", "role"}
    if not required_fields.issubset(employee.keys()):
        missing = required_fields - employee.keys()
        return {"error": f"Missing fields: {', '.join(missing)}"}
    
    if any(emp["id"] == employee["id"] for emp in EMPLOYEES_DATA["employees"]):
        return {"error": "Employee with this ID already exists."}
    
    EMPLOYEES_DATA["employees"].append(employee)
    
    # Save updated data to JSON file
    try:
        with open(EMPLOYEE_DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(EMPLOYEES_DATA, f, indent=4)
    except Exception as e:
        return {"error": f"Failed to save employee data: {str(e)}"}
    
    return {"message": "Employee added successfully", "employee": employee}

@app.put("/employees/{employee_id}")
async def edit_employee(employee_id: int, updated_employee: Dict = Body(...)):
    """Edit an existing employee's data and save persistently"""
    for idx, emp in enumerate(EMPLOYEES_DATA["employees"]):
        if emp["id"] == employee_id:
            # Update employee data
            EMPLOYEES_DATA["employees"][idx].update(updated_employee)
            # Save updated data to JSON file
            try:
                with open(EMPLOYEE_DATA_FILE, 'w', encoding='utf-8') as f:
                    json.dump(EMPLOYEES_DATA, f, indent=4)
            except Exception as e:
                return {"error": f"Failed to save employee data: {str(e)}"}
            return {"message": "Employee updated successfully", "employee": EMPLOYEES_DATA["employees"][idx]}
    return {"error": "Employee not found"}

@app.delete("/employees/{employee_id}")
async def delete_employee(employee_id: int):
    """Delete an employee by ID and save persistently"""
    for idx, emp in enumerate(EMPLOYEES_DATA["employees"]):
        if emp["id"] == employee_id:
            EMPLOYEES_DATA["employees"].pop(idx)
            try:
                with open(EMPLOYEE_DATA_FILE, 'w', encoding='utf-8') as f:
                    json.dump(EMPLOYEES_DATA, f, indent=4)
            except Exception as e:
                return {"error": f"Failed to save employee data: {str(e)}"}
            return {"message": "Employee deleted successfully"}
    return {"error": "Employee not found"}

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None

@app.post("/match-jd-resumes")
async def match_jd_resumes(
    jd_text: str = Form(None),
    jd_file: UploadFile = File(None),
    resumes: List[UploadFile] = File(...)
):
    import traceback
    try:
        # Extract JD text
        jd_content = jd_text or ""
        if jd_file is not None:
            content = await jd_file.read()
            if jd_file.filename.endswith('.pdf') and PyPDF2:
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                jd_content = " ".join(page.extract_text() or '' for page in reader.pages)
            elif jd_file.filename.endswith('.docx') and docx:
                doc = docx.Document(io.BytesIO(content))
                jd_content = " ".join([p.text for p in doc.paragraphs])
            elif jd_file.filename.endswith('.txt'):
                jd_content = content.decode(errors='ignore')
        jd_content = (jd_content or '').strip()
        if not jd_content or len(jd_content) < 20:
            print("[JD ERROR] No valid JD text provided.")
            raise HTTPException(status_code=400, detail="No valid JD text provided.")

        # Extract resume texts
        resume_texts = []
        for resume in resumes:
            content = await resume.read()
            text = ""
            if resume.filename.endswith('.pdf') and PyPDF2:
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = " ".join(page.extract_text() or '' for page in reader.pages)
            elif resume.filename.endswith('.docx') and docx:
                doc = docx.Document(io.BytesIO(content))
                text = " ".join([p.text for p in doc.paragraphs])
            elif resume.filename.endswith('.txt'):
                text = content.decode(errors='ignore')
            text = (text or '').strip()
            if len(text) < 20:
                print(f"[RESUME ERROR] Resume {resume.filename} too short or empty.")
                continue  # Ignore empty/very short resumes
            # Improved info extraction
            import re
            name = None
            email = None
            contact = None
            role = None
            department = None
            skills = []
            experience_years = None
            availability = None
            # Extract email
            email_match = re.search(r'[\w\.-]+@[\w\.-]+', text)
            if email_match:
                email = email_match.group(0)
            # Extract contact (simple phone pattern)
            contact_match = re.search(r'(\+?\d[\d\s\-]{7,}\d)', text)
            if contact_match:
                contact = contact_match.group(0)
            # Try to extract name (first line or 'Name: ...')
            lines = text.splitlines()
            for line in lines[:10]:
                if 'name:' in line.lower():
                    name = line.split(':', 1)[-1].strip()
                elif not name and len(line.split()) <= 5 and line.strip():
                    name = line.strip()
                if 'role:' in line.lower():
                    role = line.split(':', 1)[-1].strip()
                if 'department:' in line.lower():
                    department = line.split(':', 1)[-1].strip()
                if 'skills:' in line.lower():
                    skills = [s.strip() for s in line.split(':', 1)[-1].split(',') if s.strip()]
                if 'experience' in line.lower():
                    match = re.search(r'(\d+)\s*years', line.lower())
                    if match:
                        experience_years = int(match.group(1))
                if 'available' in line.lower():
                    availability = 'available'
                if 'busy' in line.lower():
                    availability = 'busy'
            # If no skills found, try to extract from text using a skill list
            if not skills:
                skill_keywords = [
                    'python', 'javascript', 'react', 'node.js', 'aws', 'docker', 'kubernetes',
                    'java', 'spring boot', 'postgresql', 'mongodb', 'sql', 'tensorflow',
                    'pytorch', 'machine learning', 'ml', 'ai', 'flutter', 'react native',
                    'vue.js', 'angular', 'django', 'fastapi', 'express', 'redis', 'kafka',
                    'go', 'rust', 'typescript', 'graphql', 'rest api', 'microservices'
                ]
                found_skills = set()
                for skill in skill_keywords:
                    if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
                        found_skills.add(skill)
                skills = list(found_skills)
            # Find matched skills with JD
            jd_skills = []
            if jd_content:
                for skill in skills:
                    if re.search(r'\b' + re.escape(skill) + r'\b', jd_content, re.IGNORECASE):
                        jd_skills.append(skill)
            resume_texts.append({
                "filename": resume.filename,
                "text": text,
                "preview": text[:300] + ("..." if len(text) > 300 else ""),
                "name": name,
                "email": email,
                "contact": contact,
                "role": role,
                "department": department,
                "skills": skills,
                "matched_skills": jd_skills,
                "experience_years": experience_years,
                "availability": availability
            })

        if not resume_texts:
            print("[RESUME ERROR] No valid resume text extracted.")
            raise HTTPException(status_code=400, detail="No valid resume text extracted.")

        # Vectorize JD and resumes
        all_texts = [jd_content] + [r['text'] for r in resume_texts]
        vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
        embeddings = vectorizer.fit_transform(all_texts)
        jd_vec = embeddings[0]
        resume_vecs = embeddings[1:]
        similarities = cosine_similarity(jd_vec, resume_vecs).flatten()

        # Attach match score to resumes and sort
        for i, sim in enumerate(similarities):
            resume_texts[i]['match_score'] = float(sim)
        matches = sorted(resume_texts, key=lambda r: r['match_score'], reverse=True)
        # Only return top 5 with score > 0.1
        matches = [m for m in matches if m['match_score'] > 0.1][:5]
        print(f"[MATCHES] Found {len(matches)} matches.")

        # --- Match against existing employees ---
        employee_matches = []
        if EMPLOYEES_DATA and EMPLOYEES_DATA.get('employees'):
            employee_texts = []
            for emp in EMPLOYEES_DATA['employees']:
                text = f"{emp.get('name','')} {emp.get('role','')} {' '.join(emp.get('skills', []))} {' '.join(emp.get('projects', []))} {emp.get('department','')} {emp.get('availability','')} {emp.get('experience_years','')} years experience"
                employee_texts.append(text.lower())
            emp_vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
            emp_embeddings = emp_vectorizer.fit_transform([jd_content] + employee_texts)
            emp_similarities = cosine_similarity(emp_embeddings[0], emp_embeddings[1:]).flatten()
            for idx, sim in enumerate(emp_similarities):
                emp = EMPLOYEES_DATA['employees'][idx].copy()
                emp['match_score'] = float(sim)
                # Find matched skills with JD
                jd_skills = []
                for skill in emp.get('skills', []):
                    if re.search(r'\b' + re.escape(skill) + r'\b', jd_content, re.IGNORECASE):
                        jd_skills.append(skill)
                emp['matched_skills'] = jd_skills
                emp['source'] = 'employee'
                employee_matches.append(emp)
        # --- Mark resume matches with source ---
        for r in resume_texts:
            r['source'] = 'resume'
        # --- Combine and sort all matches ---
        all_matches = resume_texts + employee_matches
        all_matches = [m for m in all_matches if m.get('match_score', 0) > 0.1]
        all_matches = sorted(all_matches, key=lambda r: r['match_score'], reverse=True)[:5]
        print(f"[MATCHES] Found {len(all_matches)} matches (resumes + employees).")
        return {"matches": all_matches}
    except Exception as e:
        print("[MATCH-JD-RESUMES ERROR]", str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error matching candidates: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
